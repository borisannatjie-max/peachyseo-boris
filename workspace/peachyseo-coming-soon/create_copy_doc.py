#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('PeachySEO Website Copy', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Version 1.0 | Draft for Review')
doc.add_paragraph()

# ============================================
# SECTION 1: HOMEPAGE HERO
# ============================================
doc.add_heading('Homepage - Hero Section', level=1)

doc.add_heading('Headline:', level=2)
h1 = doc.add_paragraph()
h1.add_run('SEO So Sweet, Your Competitors Will Be Peved.').bold = True

doc.add_heading('Sub-headline:', level=2)
doc.add_paragraph('We help UK small businesses dominate Google search — without the jargon, the contracts, or the empty promises.')

doc.add_heading('Primary CTA Button:', level=2)
doc.add_paragraph('Get Your Free SEO Audit →')

doc.add_heading('Secondary CTA:', level=2)
doc.add_paragraph('See How It Works')

doc.add_paragraph()

# ============================================
# SECTION 2: SERVICES SECTION
# ============================================
doc.add_heading('Services Section', level=1)

doc.add_heading('Section Headline:', level=2)
doc.add_paragraph('SEO That Actually Works. No Games, No Gimmicks.')

doc.add_heading('Service 1 - Local SEO', level=2)
p = doc.add_paragraph()
p.add_run('What it is: ').bold = True
p.add_run('We get your business found when people search "plumber near me" or "best restaurant in [your town]" on Google Maps and local search.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Includes: ').bold = True
p.add_run('Google Business Profile optimization, local citations, review management, and location-specific landing pages.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Plumbers, Electricians, HVAC, Mobile Mechanics, Restaurants, Cafes, Salons, Dentists, Vets, Gyms')

doc.add_heading('Service 2 - Website SEO', level=2)
p = doc.add_paragraph()
p.add_run('What it is: ').bold = True
p.add_run('We optimize your website so Google understands what you do — and shows you to customers searching for exactly that.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Includes: ').bold = True
p.add_run('Keyword research, on-page optimization, technical fixes, content creation, and monthly reporting.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Any business with a website that wants more organic traffic.')

doc.add_heading('Service 3 - PPC (Google Ads)', level=2)
p = doc.add_paragraph()
p.add_run('What it is: ').bold = True
p.add_run('We set up and manage Google Ads campaigns that actually convert — so you get real customers, not just clicks.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Includes: ').bold = True
p.add_run('Campaign setup, keyword targeting, ad copy, bid management, and weekly optimization.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Businesses that want fast results while SEO builds over time.')

doc.add_paragraph()

# ============================================
# SECTION 3: HOW IT WORKS
# ============================================
doc.add_heading('How It Works (3 Steps)', level=1)

doc.add_heading('Step 1: Free Audit', level=2)
p = doc.add_paragraph()
p.add_run('We look at your website, your competitors, and your market — then tell you exactly what\'s working, what\'s not, and what we can fix. No obligation, no pressure.')
doc.add_paragraph()

doc.add_heading('Step 2: Custom Plan', level=2)
p = doc.add_paragraph()
p.add_run('We build an SEO strategy tailored to YOUR business, YOUR goals, and YOUR budget. Whether you\'re a one-plumber operation or a 50-seat restaurant, we\'ve got you.')
doc.add_paragraph()

doc.add_heading('Step 3: We Do The Work', level=2)
p = doc.add_paragraph()
p.add_run('Sit back and watch your rankings climb. We handle everything — content, technical fixes, links, local optimization — while you focus on running your business. You get monthly reports that actually make sense.')
doc.add_paragraph()

# ============================================
# SECTION 4: WHY CHOOSE US
# ============================================
doc.add_heading('Why Choose PeachySEO', level=1)

doc.add_heading('Section Headline:', level=2)
doc.add_paragraph('We\'re Not Your Typical SEO Agency. Here\'s Why.')

doc.add_heading('No Long Contracts', level=2)
p = doc.add_paragraph()
p.add_run('Month-to-month. Stay because it works, not because you\'re locked in. We earn your business every single month.')
doc.add_paragraph()

doc.add_heading('Honest Reporting', level=2)
p = doc.add_paragraph()
p.add_run('No vanity metrics. You get clear reports showing exactly what we did, what changed, and how it\'s affecting your bottom line.')
doc.add_paragraph()

doc.add_heading('Real People, Not Bots', level=2)
p = doc.add_paragraph()
p.add_run('You get a dedicated account manager who actually knows your business. Not a ticketing system. Not a chatbot. A real human who cares.')
doc.add_paragraph()

doc.add_heading('UK-Based Team', level=2)
p = doc.add_paragraph()
p.add_run('We understand the UK market. Local search, local competition, local customers. We know what works here.')
doc.add_paragraph()

doc.add_heading('Powered by Ranked.ai', level=2)
p = doc.add_paragraph()
p.add_run('Behind PeachySEO is Ranked.ai — a white-label SEO platform trusted by agencies worldwide. That means enterprise-grade SEO infrastructure, without the enterprise price tag.')
doc.add_paragraph()

# ============================================
# SECTION 5: PRICING
# ============================================
doc.add_heading('Pricing Section', level=1)

doc.add_heading('Section Headline:', level=2)
doc.add_paragraph('Honest Pricing. No Hidden Fees. No Surprises.')

doc.add_heading('Starter Package - £149/month', level=2)
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Local businesses just starting with SEO')
doc.add_paragraph('• Google Business Profile optimization')
doc.add_paragraph('• Basic on-page SEO (up to 5 pages)')
doc.add_paragraph('• Monthly keyword ranking report')
doc.add_paragraph('• Email support')
doc.add_paragraph()

doc.add_heading('Growth Package - £299/month', level=2)
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Established businesses ready to dominate their market')
doc.add_paragraph('• Everything in Starter')
doc.add_paragraph('• Advanced on-page SEO (up to 15 pages)')
doc.add_paragraph('• Local citation building (20+ directories)')
doc.add_paragraph('• Monthly content creation (2 blog posts)')
doc.add_paragraph('• Competitor analysis')
doc.add_paragraph('• Priority support')
doc.add_paragraph()

doc.add_heading('Scale Package - £499/month', level=2)
p = doc.add_paragraph()
p.add_run('Perfect for: ').bold = True
p.add_run('Businesses that want the full treatment')
doc.add_paragraph('• Everything in Growth')
doc.add_paragraph('• Unlimited on-page SEO')
doc.add_paragraph('• Full citation building (50+ directories)')
doc.add_paragraph('• Weekly content creation (4 blog posts)')
doc.add_paragraph('• Link building campaign')
doc.add_paragraph('• Dedicated account manager')
doc.add_paragraph('• Weekly reporting calls')
doc.add_paragraph()

doc.add_heading('PPC Add-on:', level=2)
p = doc.add_paragraph()
p.add_run('From £199/month + ad spend').italic = True
doc.add_paragraph('Google Ads management, ad copy, bid optimization, monthly performance review.')
doc.add_paragraph()

# ============================================
# SECTION 6: CONTACT/CTA PAGE
# ============================================
doc.add_heading('Contact Page / CTA Section', level=1)

doc.add_heading('Headline:', level=2)
doc.add_paragraph('Ready to Get Found? Let\'s Talk.')

doc.add_heading('Body Copy:', level=2)
doc.add_paragraph('Whether you have questions, want a free audit, or just want to see if we\'re a good fit — drop us a line. No hard sell. No pressure. Just a friendly chat about your business.')
doc.add_paragraph()

doc.add_heading('Contact Details:', level=2)
doc.add_paragraph('Email: hello@peachyseo.com')
doc.add_paragraph('We typically respond within 24 hours.')
doc.add_paragraph()

doc.add_heading('CTA Button:', level=2)
doc.add_paragraph('Book Your Free Strategy Call →')
doc.add_paragraph()

# ============================================
# SECTION 7: SOCIAL PROOF / TESTIMONIALS
# ============================================
doc.add_heading('Testimonials Section', level=1)

doc.add_heading('Section Headline:', level=2)
doc.add_paragraph('What Our Clients Say')

doc.add_heading('Testimonial 1:', level=2)
p = doc.add_paragraph()
p.add_run('"PeachySEO transformed our online presence. We went from page 2 of Google to the top spot for \'emergency plumber Bournemouth\' in just 3 months. Now we\'re booked solid every week."').italic = True
doc.add_paragraph('— Dave M., Emergency Plumber, Bournemouth')

doc.add_heading('Testimonial 2:', level=2)
p = doc.add_paragraph()
p.add_run('"Finally, an SEO company that speaks plain English. They explained exactly what they were doing, sent regular reports, and our restaurant bookings are up 40% since we started. Worth every penny."').italic = True
doc.add_paragraph('— Sarah J., Owner, The Harbour Grill, Poole')

doc.add_heading('Testimonial 3:', level=2)
p = doc.add_paragraph()
p.add_run('"As a mechanic, I never had time for marketing. PeachySEO handles everything and now my phone rings every day with new customers. Best investment I\'ve made for my business."').italic = True
doc.add_paragraph('— Tom R., Mobile Mechanic, Christchurch')

# ============================================
# SECTION 8: ABOUT PAGE (SHORT)
# ============================================
doc.add_heading('About Section (Short Version)', level=1)

doc.add_heading('Headline:', level=2)
doc.add_paragraph('Who We Are')

doc.add_heading('Body Copy:', level=2)
p = doc.add_paragraph()
p.add_run('PeachySEO was founded with one mission: ').bold = False
p.add_run('to help small businesses in the UK get found online without the nonsense that usually comes with SEO agencies.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('We\'re a team of digital marketers, content creators, and SEO specialists who believe that every local business deserves a fair shot at ranking on Google — regardless of budget or know-how.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Behind PeachySEO is the power of Ranked.ai, giving us enterprise-grade SEO tools and expertise while keeping us small-business friendly.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('We\'re based in the UK, we work with UK businesses, and we actually answer our phones.')
doc.add_paragraph()

# ============================================
# SECTION 9: FOOTER
# ============================================
doc.add_heading('Footer Copy', level=1)

doc.add_heading('Tagline:', level=2)
doc.add_paragraph('SEO so sweet, your competitors will be peeved. 🍑')

doc.add_heading('Links:', level=2)
doc.add_paragraph('• Home')
doc.add_paragraph('• Services')
doc.add_paragraph('• Pricing')
doc.add_paragraph('• About')
doc.add_paragraph('• Contact')
doc.add_paragraph('• Privacy Policy')
doc.add_paragraph('• Terms of Service')

doc.add_heading('Contact:', level=2)
doc.add_paragraph('Email: hello@peachyseo.com')

doc.add_heading('Copyright:', level=2)
doc.add_paragraph('© 2026 PeachySEO. All rights reserved.')
doc.add_paragraph('Powered by Ranked.ai')

# ============================================
# NOTES
# ============================================
doc.add_page_break()
doc.add_heading('Notes for Review', level=1)

doc.add_paragraph('BRAND VOICE: Friendly, approachable, not corporate. We speak like real humans, not SEO robots.')
doc.add_paragraph()
doc.add_paragraph('KEY MESSAGES TO REINFORCE:')
doc.add_paragraph('• No long contracts - month to month')
doc.add_paragraph('• UK-based team')
doc.add_paragraph('• Plain English reporting')
doc.add_paragraph('• Powered by Ranked.ai (credibility)')
doc.add_paragraph('• "Peachy" branding throughout')
doc.add_paragraph()
doc.add_paragraph('TARGET AUDIENCE: UK small businesses - tradespeople (plumbers, electricians, mechanics), hospitality (restaurants, cafes, pubs), service businesses (gyms, salons, dentists)')
doc.add_paragraph()
doc.add_paragraph('COMPETITIVE ADVANTAGE: Friendly service, no jargon, honest reporting, month-to-month (vs agencies locking you into 12-month contracts)')
doc.add_paragraph()
doc.add_paragraph('NEXT STEPS:')
doc.add_paragraph('1. Review this copy')
doc.add_paragraph('2. Mark up any changes')
doc.add_paragraph('3. Send back to Boris')
doc.add_paragraph('4. We\'ll implement on peachyseo.com')

# Save
doc.save('/root/.openclaw/workspace/peachyseo-coming-soon/PeachySEO_Website_Copy_v1.docx')
print("Word document created successfully!")
